import os
import uuid
import asyncio
import sys
import time
import zipfile
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
import threading
import pdfplumber
from docx import Document
from markitdown import MarkItDown
from langchain_experimental.text_splitter import SemanticChunker
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import requests
import json

class NomicEmbedTextV2MoeEmbeddings(Embeddings):
    def __init__(self, model_path: str, device: str = "auto", max_length: int = 2048, normalize: bool = True):
        import torch
        from transformers import AutoModel, AutoTokenizer

        self.model_path = model_path
        self.max_length = max_length
        self.normalize = normalize
        resolved_device = device
        if (device or "").lower() == "auto":
            resolved_device = "cuda" if torch.cuda.is_available() else "cpu"
        self.device = torch.device(resolved_device)
        self.tokenizer = AutoTokenizer.from_pretrained(model_path, trust_remote_code=True, local_files_only=True)
        self.model = AutoModel.from_pretrained(model_path, trust_remote_code=True, local_files_only=True)
        self.model.to(self.device)
        self.model.eval()

    def _embed_batch(self, texts: List[str]) -> List[List[float]]:
        import torch

        encoded = self.tokenizer(
            texts,
            padding=True,
            truncation=True,
            max_length=self.max_length,
            return_tensors="pt",
        )
        encoded = {k: v.to(self.device) for k, v in encoded.items()}
        with torch.no_grad():
            out = self.model(**encoded)
            last_hidden = out.last_hidden_state
            mask = encoded["attention_mask"].unsqueeze(-1).type_as(last_hidden)
            summed = (last_hidden * mask).sum(dim=1)
            denom = mask.sum(dim=1).clamp(min=1e-9)
            embeddings = summed / denom
            if self.normalize:
                embeddings = torch.nn.functional.normalize(embeddings, p=2, dim=1)
        return embeddings.cpu().tolist()

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        batch_size = 16
        outputs: List[List[float]] = []
        for i in range(0, len(texts), batch_size):
            outputs.extend(self._embed_batch(texts[i : i + batch_size]))
        return outputs

    def embed_query(self, text: str) -> List[float]:
        return self._embed_batch([text])[0]

class DocumentService:
    def __init__(self):
        app_dir = os.path.dirname(os.path.dirname(__file__))
        self._model_path = os.path.join(app_dir, "models", "nomic-embed-text-v2-moe")
        self._embed_device = (os.environ.get("EMBEDDINGS_DEVICE") or "auto").strip() or "auto"
        self._chunker_lock = threading.Lock()
        self.embeddings: Optional[NomicEmbedTextV2MoeEmbeddings] = None
        self.text_splitter: Optional[SemanticChunker] = None
        self.fallback_splitter = RecursiveCharacterTextSplitter(
            chunk_size=650,
            chunk_overlap=120,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", ";", "；", ",", "，", " ", ""],
        )
        self.markitdown = MarkItDown()
        self.logger = logging.getLogger("app.document_service")
        # MinerU API配置
        self.mineru_token = os.environ.get("MINERU_API_TOKEN", "eyJ0eXBlIjoiSldUIiwiYWxnIjoiSFM1MTIifQ.eyJqdGkiOiI5NjAwMDE0NyIsInJvbCI6IlJPTEVfUkVHSVNURVIiLCJpc3MiOiJPcGVuWExhYiIsImlhdCI6MTc3MjAwODUzNSwiY2xpZW50SWQiOiJsa3pkeDU3bnZ5MjJqa3BxOXgydyIsInBob25lIjoiIiwib3BlbklkIjpudWxsLCJ1dWlkIjoiNDM2MGJlNWYtODlkNS00Mzg2LTk2NGUtYzAzOWU1ZTBlMWEzIiwiZW1haWwiOiIiLCJleHAiOjE3Nzk3ODQ1MzV9.5Cb4jvRV1HYHCqPqQFrt0WqPmBsOQxMQtIDkKId-i9bz0rLeQTLnT9fzF-EpuBIxYiHcLeduM9R_5I2SxMBSUA")
        self.mineru_api_url = "https://mineru.net/api/v4/file-urls/batch"
        self.mineru_result_url = "https://mineru.net/api/v4/extract-results/batch/{batch_id}"

    def _looks_like_pdf_password_error(self, exc: Exception) -> bool:
        cur: Optional[BaseException] = exc
        for _ in range(8):
            if cur is None:
                break
            name = cur.__class__.__name__
            if name in {"PDFPasswordIncorrect", "PDFEncryptionError"}:
                return True
            cur = getattr(cur, "__cause__", None) or getattr(cur, "__context__", None)
        msg = (str(exc) or "").lower()
        rep = (repr(exc) or "").lower()
        if ("pdfpasswordincorrect" in msg) or ("pdfpasswordincorrect" in rep):
            return True
        if ("encrypted" in msg and "password" in msg) or ("incorrect password" in msg) or ("password error" in msg):
            return True
        return False

    def _get_upload_urls(self, file_name: str, token: Optional[str] = None) -> Optional[Dict[str, Any]]:
        """获取文件上传URL"""
        try:
            effective_token = token or self.mineru_token
            header = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {effective_token}"
            }
            data = {
                "files": [
                    {"name": file_name, "data_id": str(uuid.uuid4())}
                ],
                "model_version": "vlm",
                "callback": "http://127.0.0.1:8000/callback",
                "seed": "your_random_seed",
                "extra_formats": ["html", "docx"]
            }
            response = requests.post(self.mineru_api_url, headers=header, json=data)
            if response.status_code == 200:
                result = response.json()
                if result.get("code") == 0:
                    return result.get("data")
                else:
                    self.logger.error("MinerU API error: %s", result.get("msg", "Unknown error"))
            else:
                self.logger.error("MinerU API request failed: %s, %s", response.status_code, response.text)
        except Exception as e:
            self.logger.exception("Error getting upload URLs: %s", str(e))
        return None

    def _upload_file(self, upload_url: str, file_path: str) -> bool:
        """上传文件"""
        try:
            with open(file_path, 'rb') as f:
                response = requests.put(upload_url, data=f)
                if response.status_code == 200:
                    return True
                else:
                    self.logger.error("File upload failed: %s, %s", response.status_code, response.text)
        except Exception as e:
            self.logger.exception("Error uploading file: %s", str(e))
        return False

    def _get_task_results(self, batch_id: str, token: Optional[str] = None) -> Optional[Dict[str, Any]]:
        """获取任务结果（带轮询）"""
        max_retries = 30
        retry_interval = 10  # 秒
        effective_token = token or self.mineru_token
        
        for retry in range(max_retries):
            try:
                url = self.mineru_result_url.format(batch_id=batch_id)
                header = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {effective_token}"
                }
                response = requests.get(url, headers=header)
                if response.status_code == 200:
                    result = response.json()
                    if result.get("code") == 0:
                        task_data = result.get('data')
                        if task_data and 'extract_result' in task_data:
                            for item in task_data['extract_result']:
                                state = item.get('state')
                                if state == 'completed' or state == 'done':
                                    return task_data
                                elif state == 'failed':
                                    self.logger.error("Task failed: %s", item.get('err_msg', 'Unknown error'))
                                    return None
                self.logger.info("Task still processing, retrying in %s seconds... (%s/%s)", retry_interval, retry+1, max_retries)
                time.sleep(retry_interval)
            except Exception as e:
                self.logger.exception("Error getting task results: %s", str(e))
                time.sleep(retry_interval)
        
        self.logger.error("Max retries reached. Task may still be processing.")
        return None

    def _download_and_extract_zip(self, zip_url: str, output_dir: str) -> bool:
        """下载并提取压缩包"""
        try:
            os.makedirs(output_dir, exist_ok=True)
            zip_file = os.path.join(output_dir, "result.zip")
            
            # 下载压缩包
            response = requests.get(zip_url, stream=True)
            with open(zip_file, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            # 提取压缩包
            with zipfile.ZipFile(zip_file, 'r') as zf:
                zf.extractall(output_dir)
            
            return True
        except Exception as e:
            self.logger.exception("Error downloading and extracting zip: %s", str(e))
            return False

    def _parse_content_list(self, layout_path: str) -> str:
        """解析content_list.json文件，提取文本内容"""
        try:
            with open(layout_path_path, 'r', encoding='utf-8') as f:
                layout = json.load(f)
            
            text_parts = []
            for item in layout:
                if item.get('type') == 'text':
                    text = item.get('text', '')
                    if text.strip():
                        text_parts.append(text)
                elif item.get('type') == 'list':
                    list_items = item.get('list_items', [])
                    for list_item in list_items:
                        if list_item.strip():
                            text_parts.append(f"- {list_item}")
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            self.logger.exception("Error parsing content list: %s", str(e))
            return ""

    def _read_pdf_bytes_for_mineru(self, file_path: str, password: Optional[str]) -> bytes:
        pwd = (password or "").strip()
        try:
            import io
            import pikepdf

            if pwd:
                try:
                    with pikepdf.open(file_path, password=pwd) as pdf:
                        buf = io.BytesIO()
                        pdf.save(buf)
                        return buf.getvalue()
                except Exception as e:
                    msg = (str(e) or "").lower()
                    if "password" in msg or "encrypted" in msg:
                        raise ValueError("pdf_password_incorrect")
            else:
                try:
                    with pikepdf.open(file_path) as pdf:
                        if getattr(pdf, "is_encrypted", False):
                            raise ValueError("pdf_password_required")
                except Exception as e:
                    msg = (str(e) or "").lower()
                    if "password" in msg or "encrypted" in msg:
                        raise ValueError("pdf_password_required")
        except ImportError:
            pass

        try:
            from mineru.cli.common import read_fn
        except Exception:
            with open(file_path, "rb") as f:
                return f.read()

        return read_fn(file_path)

    def _try_mineru_parse(self, file_path: str, password: Optional[str], token: Optional[str] = None) -> Optional[str]:
        """使用MinerU API解析文档"""
        try:
            if password and password.strip():
                self.logger.warning("MinerU API does not support password-protected files")
                return None

            file_name = os.path.basename(file_path)
            self.logger.info("mineru_api_parse_start file=%s", file_name)

            # 获取上传URL
            upload_data = self._get_upload_urls(file_name, token=token)
            if not upload_data:
                self.logger.error("Failed to get upload URLs")
                return None

            batch_id = upload_data.get("batch_id")
            upload_urls = upload_data.get("file_urls", [])
            if not batch_id or not upload_urls:
                self.logger.error("Invalid upload data: batch_id=%s, urls=%s", batch_id, upload_urls)
                return None

            # 上传文件
            if not self._upload_file(upload_urls[0], file_path):
                self.logger.error("Failed to upload file")
                return None

            # 获取任务结果
            task_data = self._get_task_results(batch_id, token=token)
            if not task_data:
                self.logger.error("Failed to get task results")
                return None

            # 下载并提取压缩包
            extract_result = task_data.get("extract_result", [])
            if not extract_result:
                self.logger.error("No extract result found")
                return None

            item = extract_result[0]
            zip_url = item.get("full_zip_url")
            if not zip_url:
                self.logger.error("No full_zip_url found")
                return None

            # 创建临时目录
            import tempfile
            import shutil
            
            # 先创建临时目录，以便后续手动清理
            tmp_dir = tempfile.mkdtemp(prefix="mineru_api_")
            try:
                # 下载并提取压缩包
                if not self._download_and_extract_zip(zip_url, tmp_dir):
                    self.logger.error("Failed to download and extract zip")
                    return None

                # 查找full.md文件
                full_md_path = None
                for root, dirs, files in os.walk(tmp_dir):
                    for file in files:
                        if file == "full.md":
                            full_md_path = os.path.join(root, file)
                            break
                    if full_md_path:
                        break

                if full_md_path:
                    # 读取full.md文件
                    try:
                        with open(full_md_path, 'r', encoding='utf-8', errors='ignore') as f:
                            text = f.read()
                        self.logger.info("Using full.md file for text extraction")
                    except Exception as e:
                        self.logger.exception("Error reading full.md: %s", str(e))
                        text = ""
                else:
                    # 查找content_list.json文件作为备选
                    content_list_path = None
                    for root, dirs, files in os.walk(tmp_dir):
                        for file in files:
                            if file == "content_list.json":
                                content_list_path = os.path.join(root, file)
                                break
                        if content_list_path:
                            break

                    if content_list_path:
                        # 解析content_list.json文件
                        text = self._parse_content_list(content_list_path)
                    else:
                        self.logger.error("Neither full.md nor content_list.json found")
                        return None
            finally:
                # 清理临时目录
                try:
                    shutil.rmtree(tmp_dir)
                    self.logger.info("Temporary directory cleaned: %s", tmp_dir)
                except Exception as e:
                    self.logger.exception("Error cleaning temporary directory: %s", str(e))
            
            if text and text.strip():
                self.logger.info("mineru_api_parse_done chars=%s", len(text))
                return text
            else:
                self.logger.warning("No text extracted from MinerU API")
                return None

        except ValueError:
            raise
        except Exception as e:
            if self._looks_like_pdf_password_error(e):
                pwd = (password or "").strip()
                raise ValueError("pdf_password_incorrect" if pwd else "pdf_password_required")
            self.logger.warning(
                "mineru_api_parse_failed err_type=%s err=%s",
                e.__class__.__name__,
                str(e)[:300],
                exc_info=True,
            )
            return None

    def _ensure_semantic_chunker(self) -> Optional[SemanticChunker]:
        if self.text_splitter is not None:
            return self.text_splitter
        with self._chunker_lock:
            if self.text_splitter is not None:
                return self.text_splitter
            try:
                self.embeddings = NomicEmbedTextV2MoeEmbeddings(model_path=self._model_path, device=self._embed_device, normalize=True)
                self.text_splitter = SemanticChunker(self.embeddings)
                return self.text_splitter
            except Exception:
                self.logger.exception("semantic_chunker_init_failed")
                self.embeddings = None
                self.text_splitter = None
                return None

    async def parse_document(self, file_path: str, file_type: str, password: Optional[str] = None, parser: Optional[str] = None, mineru_token: Optional[str] = None) -> str:
        """Parse different file types and return plain text."""
        text = ""
        if file_type == "pdf":
            pwd = (password or "").strip()
            mode = (parser or "").strip().lower()
            if mode not in {"auto", "mineru", "fallback", "pdfplumber"}:
                mode = "auto"

            if mode in {"auto", "mineru"}:
                try:
                    mineru_text = await asyncio.to_thread(self._try_mineru_parse, file_path, pwd or None, mineru_token)
                    if mineru_text and mineru_text.strip():
                        cleaned = self._clean_markdown_for_rag(mineru_text)
                        try:
                            size = os.path.getsize(file_path)
                        except Exception:
                            size = -1
                        self.logger.info("document_parsed type=pdf parser=mineru chars=%s bytes=%s", len(cleaned), size)
                        return cleaned
                except ValueError:
                    raise
                except Exception:
                    if mode == "mineru":
                        raise
                    pass

            try:
                with pdfplumber.open(file_path, password=(pwd or None)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                if self._looks_like_pdf_password_error(e):
                    raise ValueError("pdf_password_incorrect" if pwd else "pdf_password_required")
                raise
        elif file_type == "docx":
            try:
                doc = Document(file_path)
                for para in doc.paragraphs:
                    if para.text:
                        text += para.text + "\n"

                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join([cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()])
                        if row_text:
                            text += row_text + "\n"
            except Exception:
                self.logger.exception("docx_parse_failed path=%s", file_path)
                text = ""
        elif file_type in ["txt", "md"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                    text = f.read()
        else:
            # Fallback to MarkItDown for other formats
            try:
                result = self.markitdown.convert(file_path)
                text = result.text_content
            except Exception:
                self.logger.exception("markitdown_convert_failed path=%s", file_path)
                text = ""
        
        cleaned = self._clean_markdown_for_rag(text) if file_type == "md" else self._clean_text(text)
        try:
            size = os.path.getsize(file_path)
        except Exception:
            size = -1
        self.logger.info("document_parsed type=%s parser=fallback chars=%s bytes=%s", file_type, len(cleaned), size)
        return cleaned

    def _clean_markdown_for_rag(self, text: str) -> str:
        import re

        s = (text or "").replace("\r\n", "\n").replace("\r", "\n")

        s = re.sub(r"<think>[\s\S]*?</think>", "", s, flags=re.IGNORECASE)
        s = re.sub(r"<analysis>[\s\S]*?</analysis>", "", s, flags=re.IGNORECASE)

        s = re.sub(r"```[^\n]*\n([\s\S]*?)```", r"\1", s)
        s = re.sub(r"`([^`\n]+)`", r"\1", s)

        s = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", s)
        s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)

        lines = s.split("\n")
        out_lines = []
        for line in lines:
            raw = line.rstrip()
            if not raw.strip():
                out_lines.append("")
                continue

            raw = re.sub(r"^\s{0,3}#{1,6}\s*", "", raw)
            raw = re.sub(r"^\s{0,3}>\s?", "", raw)
            raw = re.sub(r"^\s*[-*+]\s+", "- ", raw)
            raw = re.sub(r"^\s*\d+\.\s+", "- ", raw)

            if "|" in raw and not re.fullmatch(r"\s*\|?[\s:-]+\|?(\s*\|[\s:-]+\|?)*\s*", raw):
                cells = [c.strip() for c in raw.strip().strip("|").split("|")]
                cells = [c for c in cells if c]
                if len(cells) >= 2:
                    raw = "\t".join(cells)

            raw = raw.replace("**", "").replace("__", "").replace("~~", "")
            raw = re.sub(r"(?<!\w)\*(?!\w)|(?<!\w)_(?!\w)", "", raw)

            out_lines.append(raw)

        return self._clean_text("\n".join(out_lines))

    def _clean_text(self, text: str) -> str:
        import re
        s = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        s = re.sub(r"[ \t\f\v]+", " ", s)
        s = re.sub(r"[ \t]+\n", "\n", s)
        s = re.sub(r"\n[ \t]+", "\n", s)
        s = re.sub(r"\n{3,}", "\n\n", s)
        return s.strip()

    def split_text(self, text: str) -> List[str]:
        """Split text using dynamic semantic chunking."""
        def _refine(chunks: List[str]) -> List[str]:
            refined: List[str] = []
            for c in chunks:
                if len(c) > 900:
                    refined.extend(self.fallback_splitter.split_text(c))
                else:
                    refined.append(c)
            return [c for c in refined if c and c.strip()]

        splitter = self._ensure_semantic_chunker()
        if splitter is not None:
            try:
                chunks = splitter.create_documents([text])
                results = [chunk.page_content for chunk in chunks if chunk.page_content and chunk.page_content.strip()]
                if results:
                    results = _refine(results)
                    self.logger.info("document_splitter=semantic chunks=%s", len(results))
                    return results
            except Exception:
                self.logger.exception("semantic_chunker_failed")

        results = self.fallback_splitter.split_text(text)
        results = [c for c in results if c and c.strip()]
        self.logger.info("document_splitter=fallback chunks=%s", len(results))
        return results

_document_service: Optional[DocumentService] = None
_document_service_lock = threading.Lock()


def get_document_service() -> DocumentService:
    global _document_service
    if _document_service is not None:
        return _document_service
    with _document_service_lock:
        if _document_service is None:
            _document_service = DocumentService()
        return _document_service
